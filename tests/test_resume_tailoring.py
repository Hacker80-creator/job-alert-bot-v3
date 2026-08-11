from __future__ import annotations

import hashlib
import json
import tempfile
import unittest
import zipfile
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import Mock, patch

from docx import Document

import job_monitor as bot
import resume_tailor as tailor


ROOT = Path(__file__).parents[1]
MASTER = ROOT / "resume" / "master_resume.docx"
EXPECTED_MASTER_HASH = "9c16cae30bae8e7a4d4f8b834a030cbdf1c2c759963407164a6538dca0505c07"


def make_job() -> SimpleNamespace:
    return SimpleNamespace(
        company="MakeMyTrip",
        title="Product Analyst",
        location="Bengaluru, India",
        url="https://careers.example.test/jobs/123456",
        requisition_id="REQ-123456",
        description=(
            "Analyze product data using Python, SQL, Exploratory Data Analysis, "
            "Statistical Analysis and Power BI. AWS, Tableau and Airflow are preferred."
        ),
    )


def make_meaningful_plan(template: tailor.TemplateSnapshot) -> dict[str, object]:
    plan = tailor.safe_plan(template, make_job())
    plan["professional_summary"] = {
        "text": (
            "Analytical DevOps professional with hands-on Python, SQL, "
            "Exploratory Data Analysis (EDA), Data Cleaning, Statistical Analysis, "
            "Power BI, CI/CD, containerization, and Machine Learning. Experienced "
            "in automated pipelines, scalable deployment workflows, ML benchmarking, "
            "and performance analysis."
        ),
        "evidence": [
            template.paragraphs[template.summary_index],
            template.paragraphs[template.skill_indices[1]],
            template.paragraphs[template.skill_indices[2]],
            template.paragraphs[template.project_indices[0]],
        ],
    }
    experience_source = template.paragraphs[template.experience_indices[3]]
    plan["experience_bullets"] = [{
        "index": 3,
        "text": (
            "Used Bash, Python, and Groovy automation scripts for build orchestration, "
            "environment validation, and log processing across Linux-based systems, "
            "automating repetitive operational workflows."
        ),
        "evidence": [experience_source],
    }]
    project_source = template.paragraphs[template.project_indices[0]]
    plan["project_bullets"] = [{
        "index": 0,
        "text": (
            "Analyzed 2,500 samples with 21 morphological features across 7 classes "
            "through a config-driven ML pipeline benchmarking 6 algorithms (SVM, "
            "Random Forest, KNN, Logistic Regression, Naive Bayes, Decision Tree)."
        ),
        "evidence": [project_source],
    }]
    plan["supported_skills"] = [
        "Python", "SQL", "Exploratory Data Analysis (EDA)",
        "Data Cleaning", "Statistical Analysis", "Power BI",
    ]
    return plan


class ResumeTailoringTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        cls.template = tailor.inspect_template(MASTER)

    def test_master_hash_and_dynamic_sections_are_stable(self) -> None:
        self.assertEqual(EXPECTED_MASTER_HASH, hashlib.sha256(MASTER.read_bytes()).hexdigest())
        self.assertEqual(EXPECTED_MASTER_HASH, self.template.sha256)
        self.assertEqual(4, len(self.template.skill_indices))
        self.assertEqual(5, len(self.template.experience_indices))
        self.assertEqual(3, len(self.template.project_indices))
        self.assertIn("Python", self.template.skill_catalog)
        self.assertIn("EDUCATION", self.template.headings)
        self.assertGreaterEqual(len(self.template.hyperlink_targets), 3)

    def test_grounded_generation_preserves_master_and_docx_structure(self) -> None:
        before = MASTER.read_bytes()
        master_doc = Document(MASTER)
        with tempfile.TemporaryDirectory() as temp_dir, patch.object(
            tailor,
            "request_tailoring_plan",
            return_value=(make_meaningful_plan(self.template), "gemini-test", ()),
        ):
            result = tailor.generate_tailored_resume(
                make_job(), output_dir=temp_dir, api_key="test"
            )
            self.assertTrue(result.path.is_file())
            self.assertEqual("gemini-test", result.model)
            self.assertTrue(result.changed_sections)
            self.assertIsNotNone(result.comparison)
            self.assertGreaterEqual(
                len(result.comparison.newly_surfaced_summary_keywords), 2
            )
            self.assertGreaterEqual(result.comparison.rewritten_bullet_count, 1)
            self.assertIn("Python", result.supported_skills)
            self.assertIn("AWS", result.important_gaps)
            tailor.validate_generated_resume(self.template, result.path)
            with zipfile.ZipFile(MASTER) as master_zip, zipfile.ZipFile(result.path) as generated_zip:
                for name in master_zip.namelist():
                    if name != "word/document.xml":
                        self.assertEqual(
                            master_zip.read(name),
                            generated_zip.read(name),
                            f"preserve-only DOCX part changed: {name}",
                        )
            generated = Document(result.path)
            self.assertNotEqual(
                [paragraph.text for paragraph in master_doc.paragraphs],
                [paragraph.text for paragraph in generated.paragraphs],
            )
            self.assertEqual(len(master_doc.paragraphs), len(generated.paragraphs))
            self.assertEqual(
                [(s.page_width, s.page_height, s.left_margin, s.right_margin) for s in master_doc.sections],
                [(s.page_width, s.page_height, s.left_margin, s.right_margin) for s in generated.sections],
            )
            for index in self.template.skill_indices:
                original_label = master_doc.paragraphs[index].runs[0]
                generated_label = generated.paragraphs[index].runs[0]
                self.assertEqual(original_label.bold, generated_label.bold)
                self.assertEqual(
                    original_label.font.color.rgb,
                    generated_label.font.color.rgb,
                )
            for index in self.template.headings.values():
                self.assertTrue(generated.paragraphs[index].paragraph_format.keep_with_next)
            self.assertIn("JAGADEV L", "\n".join(p.text for p in generated.paragraphs))
        self.assertEqual(before, MASTER.read_bytes())

    def test_unsupported_claims_are_rejected(self) -> None:
        job = make_job()
        raw = tailor.safe_plan(self.template, job)
        raw["professional_summary"] = {
            "text": "Kubernetes architect with 99% production availability.",
            "evidence": [self.template.paragraphs[self.template.summary_index]],
        }
        raw["experience_bullets"] = [{
            "index": 0,
            "text": "•  Built Kubernetes systems with 99% availability.",
            "evidence": [self.template.paragraphs[self.template.experience_indices[0]]],
        }]
        plan, warnings = tailor.validate_plan(raw, self.template, job)
        self.assertEqual(
            self.template.paragraphs[self.template.summary_index],
            plan["professional_summary"]["text"],
        )
        self.assertEqual([], plan["experience_bullets"])
        self.assertTrue(any("unsafe summary" in warning for warning in warnings))
        self.assertTrue(any("experience_bullets" in warning for warning in warnings))

    def test_product_analyst_ranking_prioritizes_data_evidence(self) -> None:
        plan = tailor.safe_plan(self.template, make_job())
        self.assertEqual(3, plan["experience_order"][0])
        self.assertEqual([0, 1, 2], plan["project_order"])
        prompt = tailor.build_prompt(self.template, make_job())
        self.assertIn("plain, natural professional English", prompt)
        self.assertIn("no more than one experience-bullet rewrite", prompt)

    def test_primary_model_falls_back_to_flash_lite(self) -> None:
        raw = tailor.safe_plan(self.template, make_job())
        with patch.object(
            tailor, "_call_model", side_effect=[RuntimeError("primary down"), raw]
        ) as call:
            plan, model, warnings = tailor.request_tailoring_plan(
                self.template, make_job(), api_key="test-key"
            )
        self.assertEqual(raw, plan)
        self.assertEqual("gemini-3.5-flash-lite", model)
        self.assertEqual(2, call.call_count)
        self.assertIn(
            "gemini-3.6-flash failed: primary down",
            warnings,
        )

    def test_reordering_only_ai_plan_is_never_attached_as_tailored(self) -> None:
        job = make_job()
        raw = tailor.safe_plan(self.template, job)
        with tempfile.TemporaryDirectory() as temp_dir, patch.object(
            tailor, "request_tailoring_plan", return_value=(raw, "gemini-test", ())
        ):
            with self.assertRaisesRegex(
                tailor.ResumeTailoringError, "substantive professional-summary"
            ):
                tailor.generate_tailored_resume(
                    job, output_dir=temp_dir, api_key="test", require_ai=True
                )

    def test_gemini_request_uses_schema_and_secret_header(self) -> None:
        raw = tailor.safe_plan(self.template, make_job())
        response = Mock()
        response.json.return_value = {
            "status": "completed",
            "steps": [{
                "type": "model_output",
                "content": [{"type": "text", "text": json.dumps(raw)}],
            }],
        }
        with patch.object(tailor.requests, "post", return_value=response) as post:
            self.assertEqual(
                raw,
                tailor._call_model("gemini-3.6-flash", "prompt", "test-secret"),
            )
        kwargs = post.call_args.kwargs
        self.assertEqual(
            "https://generativelanguage.googleapis.com/v1beta/interactions",
            post.call_args.args[0],
        )
        self.assertEqual("prompt", kwargs["json"]["input"])
        self.assertEqual("test-secret", kwargs["headers"]["x-goog-api-key"])
        self.assertNotIn("test-secret", post.call_args.args[0])
        self.assertEqual(
            "application/json",
            kwargs["json"]["response_format"]["mime_type"],
        )
        self.assertEqual(
            tailor.TAILOR_SCHEMA,
            kwargs["json"]["response_format"]["schema"],
        )

    def test_gemini_error_exposes_safe_api_message(self) -> None:
        response = Mock(
            ok=False,
            status_code=403,
        )
        response.json.return_value = {
            "error": {
                "code": "permission_denied",
                "message": "API key lacks access",
            }
        }
        with self.assertRaisesRegex(
            tailor.ResumeTailoringError,
            r"HTTP 403.*permission_denied.*lacks access",
        ):
            tailor._response_json(response)

    def test_metric_or_skill_cannot_move_between_bullets(self) -> None:
        job = make_job()
        first = self.template.paragraphs[self.template.experience_indices[0]]
        second = self.template.paragraphs[self.template.experience_indices[1]]
        raw = tailor.safe_plan(self.template, job)
        raw["experience_bullets"] = [{
            "index": 1,
            "text": "•  Containerized 20+ repositories using Docker.",
            "evidence": [second],
        }]
        plan, warnings = tailor.validate_plan(raw, self.template, job)
        self.assertIn("20+", first)
        self.assertNotIn("20+", second)
        self.assertEqual([], plan["experience_bullets"])
        self.assertTrue(any("experience_bullets[1]" in value for value in warnings))

    def test_output_name_contains_company_role_and_job_id(self) -> None:
        filename = tailor.output_filename(make_job())
        self.assertEqual("MakeMyTrip_Product_Analyst_REQ_123456_Jagadev.docx", filename)

    def test_discord_uses_multipart_and_preserves_existing_card(self) -> None:
        response = Mock(status_code=204, text="")
        job = bot.Job(
            company="MakeMyTrip",
            title="Product Analyst",
            location="Bengaluru, India",
            url="https://careers.example.test/jobs/123456",
            source="Official careers: test",
            score=90,
            reasons=["title match", "Python and SQL"],
            requisition_id="REQ-123456",
        )
        result = tailor.TailoredResume(
            path=MASTER,
            supported_skills=("Python", "SQL"),
            important_gaps=("AWS",),
            model="gemini-3.6-flash",
            changed_sections=("professional summary rewritten",),
            comparison=tailor.ATSComparison(
                summary_keywords_before=("Python",),
                summary_keywords_after=("Python", "SQL"),
                newly_surfaced_summary_keywords=("SQL",),
                experience_bullets_rewritten=1,
                project_bullets_rewritten=0,
            ),
            source_path=MASTER,
        )
        original_webhook = bot.DISCORD_WEBHOOK_URL
        try:
            bot.DISCORD_WEBHOOK_URL = "https://discord.example/webhook"
            with patch.object(bot.requests, "post", return_value=response) as post:
                self.assertTrue(
                    bot.discord_post(job, result, include_original_resume=True)
                )
            kwargs = post.call_args.kwargs
            self.assertIn("files", kwargs)
            self.assertNotIn("json", kwargs)
            payload = json.loads(kwargs["data"]["payload_json"])
            fields = payload["embeds"][0]["fields"]
            names = {field["name"] for field in fields}
            self.assertTrue(
                {"Location", "Expected salary*", "Match score", "WLB priority", "Source"}.issubset(names)
            )
            self.assertTrue(
                {
                    "Resume-supported skills", "Important gaps", "Tailored resume",
                    "ATS evidence comparison",
                }.issubset(names)
            )
            attachment = kwargs["files"]["files[0]"]
            self.assertEqual("master_resume.docx", attachment[0])
            self.assertGreater(len(attachment[1]), 1_000)
            original = kwargs["files"]["files[1]"]
            self.assertEqual("ORIGINAL_Jagadev.docx", original[0])
            self.assertGreater(len(original[1]), 1_000)
        finally:
            bot.DISCORD_WEBHOOK_URL = original_webhook

    def test_resume_worker_failure_never_suppresses_job_alert(self) -> None:
        job = bot.Job("Example", "Data Analyst", "Bengaluru", "https://example.test/jobs/1", "test")
        original_enabled = bot.ENABLE_RESUME_TAILORING
        try:
            bot.ENABLE_RESUME_TAILORING = True
            with patch.object(
                tailor, "generate_tailored_resume", side_effect=RuntimeError("boom")
            ):
                self.assertEqual({}, bot.prepare_tailored_resumes([job]))
        finally:
            bot.ENABLE_RESUME_TAILORING = original_enabled

    def test_workflow_is_branch_safe_and_configures_models(self) -> None:
        workflow = (ROOT / ".github" / "workflows" / "job-alerts.yml").read_text(encoding="utf-8")
        self.assertIn("ref: ${{ github.ref }}", workflow)
        self.assertIn("github.ref_name != 'main'", workflow)
        self.assertIn("always() && github.ref_name == 'main'", workflow)
        self.assertIn("GEMINI_API_KEY: ${{ secrets.GEMINI_API_KEY }}", workflow)
        self.assertIn('GEMINI_PRIMARY_MODEL: "gemini-3.6-flash"', workflow)
        self.assertIn('GEMINI_FALLBACK_MODEL: "gemini-3.5-flash-lite"', workflow)
        self.assertNotIn("ANTHROPIC_API_KEY", workflow)


if __name__ == "__main__":
    unittest.main()
