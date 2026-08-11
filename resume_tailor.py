"""Truthful, format-preserving resume tailoring for Discord job alerts.

Gemini proposes structured edits. Deterministic validation checks every
proposal against the immutable master before modifying a copy.
"""
from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import time
import zipfile
from collections import Counter
from copy import deepcopy
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

import requests
from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).parent
DEFAULT_MASTER_RESUME = ROOT / "resume" / "master_resume.docx"
DEFAULT_OUTPUT_DIR = ROOT / "resumes" / "generated"
GEMINI_INTERACTIONS_URL = "https://generativelanguage.googleapis.com/v1beta/interactions"
PRIMARY_MODEL = os.getenv("GEMINI_PRIMARY_MODEL", "gemini-3.6-flash").strip()
FALLBACK_MODEL = os.getenv(
    "GEMINI_FALLBACK_MODEL", "gemini-3.5-flash-lite"
).strip()
REQUIRED_HEADINGS = (
    "PROFESSIONAL SUMMARY",
    "TECHNICAL SKILLS",
    "EXPERIENCE",
    "PROJECTS",
    "EDUCATION",
)
OPTIONAL_HEADINGS = ("CERTIFICATION", "CERTIFICATIONS")
PLACEHOLDER_MARKERS = (
    "[company]", "[role]", "[job id]", "<company>", "<role>",
    "tbd", "todo", "lorem ipsum",
)
KNOWN_TECH_TERMS = (
    "airflow", "aws", "azure", "bigquery", "c++", "databricks", "dbt",
    "docker", "gcp", "git", "github", "groovy", "hadoop", "jenkins",
    "jfrog", "kafka", "keras", "kubernetes", "langchain", "linux", "llm",
    "looker", "matplotlib", "mlflow", "mongodb", "mysql", "nlp", "numpy",
    "oracle", "pandas", "podman", "postgresql", "power bi", "pytorch",
    "rag", "redshift", "scikit-learn", "seaborn", "snowflake", "spark",
    "sql", "tableau", "tensorflow", "terraform", "vector database",
)
GAP_CANDIDATES = (
    "AWS", "Azure", "GCP", "Kubernetes", "Terraform", "Spark", "Airflow",
    "Kafka", "Databricks", "Snowflake", "dbt", "Tableau", "Looker",
    "PyTorch", "TensorFlow", "LLM", "RAG", "NLP", "MLflow",
)
STOPWORDS = {
    "about", "across", "after", "again", "against", "along", "also", "among",
    "and", "are", "been", "before", "being", "between", "build", "building",
    "built", "can", "delivery", "developed", "developing", "for", "from",
    "have", "into", "more", "most", "over", "production", "professional",
    "role", "scalable", "systems", "that", "the", "their", "these", "this",
    "through", "using", "with", "work", "worked", "working",
}


TAILOR_SCHEMA: dict[str, Any] = {
    "type": "object",
    "properties": {
        "professional_summary": {
            "type": "object",
            "properties": {
                "text": {"type": "string"},
                "evidence": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["text", "evidence"],
            "additionalProperties": False,
        },
        "skill_priorities": {"type": "array", "items": {"type": "string"}},
        "experience_bullets": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer"},
                    "text": {"type": "string"},
                    "evidence": {"type": "array", "items": {"type": "string"}},
                },
                "required": ["index", "text", "evidence"],
                "additionalProperties": False,
            },
        },
        "experience_order": {"type": "array", "items": {"type": "integer"}},
        "project_bullets": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer"},
                    "text": {"type": "string"},
                    "evidence": {"type": "array", "items": {"type": "string"}},
                },
                "required": ["index", "text", "evidence"],
                "additionalProperties": False,
            },
        },
        "project_order": {"type": "array", "items": {"type": "integer"}},
        "supported_skills": {"type": "array", "items": {"type": "string"}},
        "important_gaps": {"type": "array", "items": {"type": "string"}},
    },
    "required": [
        "professional_summary", "skill_priorities", "experience_bullets",
        "experience_order", "project_bullets", "project_order",
        "supported_skills", "important_gaps",
    ],
    "additionalProperties": False,
}


class ResumeTailoringError(RuntimeError):
    """Raised when an AI-backed resume cannot be safely produced."""


@dataclass(frozen=True)
class TemplateSnapshot:
    path: Path
    sha256: str
    paragraphs: tuple[str, ...]
    headings: dict[str, int]
    summary_index: int
    skill_indices: tuple[int, ...]
    experience_indices: tuple[int, ...]
    project_indices: tuple[int, ...]
    skill_catalog: tuple[str, ...]
    package_entries: tuple[str, ...]
    hyperlink_targets: tuple[str, ...]
    section_signature: tuple[tuple[Any, ...], ...]
    header_text: tuple[str, ...]
    footer_text: tuple[str, ...]

    @property
    def full_text(self) -> str:
        return "\n".join(self.paragraphs)


@dataclass(frozen=True)
class ATSComparison:
    summary_keywords_before: tuple[str, ...]
    summary_keywords_after: tuple[str, ...]
    newly_surfaced_summary_keywords: tuple[str, ...]
    experience_bullets_rewritten: int
    project_bullets_rewritten: int

    @property
    def rewritten_bullet_count(self) -> int:
        return self.experience_bullets_rewritten + self.project_bullets_rewritten


@dataclass(frozen=True)
class TailoredResume:
    path: Path
    supported_skills: tuple[str, ...]
    important_gaps: tuple[str, ...]
    model: str
    warnings: tuple[str, ...] = ()
    changed_sections: tuple[str, ...] = ()
    comparison: ATSComparison | None = None
    source_path: Path | None = None


@dataclass(frozen=True)
class JobContext:
    company: str
    title: str
    location: str
    url: str
    description: str
    requisition_id: str = ""


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _normalize(value: str) -> str:
    return re.sub(r"[^a-z0-9+#./]+", " ", str(value).casefold()).strip()


def _contains_phrase(haystack: str, needle: str) -> bool:
    normalized_needle = _normalize(needle)
    return bool(normalized_needle) and (
        f" {normalized_needle} " in f" {_normalize(haystack)} "
    )


def _section_signature(document: Any) -> tuple[tuple[Any, ...], ...]:
    return tuple(
        (
            section.page_width, section.page_height, section.orientation,
            section.top_margin, section.bottom_margin, section.left_margin,
            section.right_margin, section.header_distance,
            section.footer_distance, section.start_type,
        )
        for section in document.sections
    )


def _package_entries(path: Path) -> tuple[str, ...]:
    with zipfile.ZipFile(path) as archive:
        return tuple(sorted(archive.namelist()))


def _preserve_only_part_hashes(path: Path) -> tuple[tuple[str, str], ...]:
    with zipfile.ZipFile(path) as archive:
        return tuple(
            (name, hashlib.sha256(archive.read(name)).hexdigest())
            for name in sorted(archive.namelist())
            if name != "word/document.xml"
        )


def _hyperlink_targets(path: Path) -> tuple[str, ...]:
    relationship_part = "word/_rels/document.xml.rels"
    with zipfile.ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read(relationship_part))
    return tuple(sorted(
        relationship.attrib.get("Target", "")
        for relationship in root
        if relationship.attrib.get("Type", "").endswith("/hyperlink")
    ))


def _find_heading_indices(paragraphs: list[Paragraph]) -> dict[str, int]:
    headings: dict[str, int] = {}
    expected = set(REQUIRED_HEADINGS + OPTIONAL_HEADINGS)
    for index, paragraph in enumerate(paragraphs):
        key = " ".join(paragraph.text.upper().split())
        if key in expected and key not in headings:
            headings[key] = index
    missing = [heading for heading in REQUIRED_HEADINGS if heading not in headings]
    if missing:
        raise ResumeTailoringError(
            "Master resume is missing required sections: " + ", ".join(missing)
        )
    return headings


def _nonempty_indices(paragraphs: list[Paragraph], start: int, end: int) -> list[int]:
    return [index for index in range(start + 1, end) if paragraphs[index].text.strip()]


def inspect_template(path: Path | str = DEFAULT_MASTER_RESUME) -> TemplateSnapshot:
    resume_path = Path(path).resolve()
    if not resume_path.is_file():
        raise ResumeTailoringError(f"Master resume not found: {resume_path}")
    resume_hash = _sha256(resume_path)
    hash_manifest = resume_path.with_suffix(".sha256")
    if hash_manifest.is_file():
        expected_hash = hash_manifest.read_text(encoding="utf-8").split()[0].casefold()
        if resume_hash.casefold() != expected_hash:
            raise ResumeTailoringError("Master resume does not match its immutable hash")
    document = Document(resume_path)
    paragraphs = list(document.paragraphs)
    headings = _find_heading_indices(paragraphs)
    summary_candidates = _nonempty_indices(
        paragraphs, headings["PROFESSIONAL SUMMARY"], headings["TECHNICAL SKILLS"]
    )
    if len(summary_candidates) != 1:
        raise ResumeTailoringError("Expected exactly one professional-summary paragraph")
    skill_indices = _nonempty_indices(
        paragraphs, headings["TECHNICAL SKILLS"], headings["EXPERIENCE"]
    )
    experience_indices = [
        index
        for index in _nonempty_indices(
            paragraphs, headings["EXPERIENCE"], headings["PROJECTS"]
        )
        if paragraphs[index].text.lstrip().startswith("•")
    ]
    project_indices = [
        index
        for index in _nonempty_indices(
            paragraphs, headings["PROJECTS"], headings["EDUCATION"]
        )
        if paragraphs[index].text.lstrip().startswith("•")
    ]
    if not skill_indices or not experience_indices or not project_indices:
        raise ResumeTailoringError("Master resume has no safely editable skills or bullets")

    skill_catalog: list[str] = []
    normalized_catalog: set[str] = set()
    for index in skill_indices:
        text = paragraphs[index].text
        if ":" not in text:
            continue
        _, values = text.split(":", 1)
        for value in values.split(","):
            skill = value.strip()
            key = _normalize(skill)
            if skill and key not in normalized_catalog:
                skill_catalog.append(skill)
                normalized_catalog.add(key)

    return TemplateSnapshot(
        path=resume_path,
        sha256=resume_hash,
        paragraphs=tuple(paragraph.text for paragraph in paragraphs),
        headings=headings,
        summary_index=summary_candidates[0],
        skill_indices=tuple(skill_indices),
        experience_indices=tuple(experience_indices),
        project_indices=tuple(project_indices),
        skill_catalog=tuple(skill_catalog),
        package_entries=_package_entries(resume_path),
        hyperlink_targets=_hyperlink_targets(resume_path),
        section_signature=_section_signature(document),
        header_text=tuple(
            "\n".join(paragraph.text for paragraph in section.header.paragraphs)
            for section in document.sections
        ),
        footer_text=tuple(
            "\n".join(paragraph.text for paragraph in section.footer.paragraphs)
            for section in document.sections
        ),
    )


def _job_context(job: Any) -> JobContext:
    return JobContext(
        company=str(getattr(job, "company", "") or ""),
        title=str(getattr(job, "title", "") or ""),
        location=str(getattr(job, "location", "") or ""),
        url=str(getattr(job, "url", "") or ""),
        description=str(getattr(job, "description", "") or ""),
        requisition_id=str(getattr(job, "requisition_id", "") or ""),
    )


def _indexed_bullets(template: TemplateSnapshot, indices: tuple[int, ...]) -> str:
    return "\n".join(
        f"[{local_index}] {template.paragraphs[paragraph_index]}"
        for local_index, paragraph_index in enumerate(indices)
    )


def build_prompt(template: TemplateSnapshot, job: Any) -> str:
    context = _job_context(job)
    description = context.description[:18_000]
    supported_skills = _skills_in_job(template, description)
    summary_skills = [
        skill
        for skill in supported_skills
        if not _contains_phrase(template.paragraphs[template.summary_index], skill)
    ]
    required_summary_skills = min(2, len(summary_skills))
    return f"""Tailor a resume conservatively for the job below.

NON-NEGOTIABLE RULES
- The master resume is the only factual source of truth.
- Never invent experience, tools, cloud platforms, metrics, dates, education, titles, certifications, or responsibilities.
- Every rewritten summary must include exact evidence quotes copied from the master resume.
- Every rewritten bullet must include exact evidence quotes copied from that same indexed source bullet.
- Never use the job description as factual evidence about the candidate.
- ATS alignment may use only truthful keywords already supported by the cited evidence.
- If a field cannot be improved safely, keep it unchanged or omit its rewrite proposal.
- Keep the summary no longer than the existing summary.
- Keep every bullet no longer than its source bullet.
- Use only skill names from ALLOWED SKILLS in skill_priorities and supported_skills.
- supported_skills are skills present in both the job description and master resume.
- important_gaps are requirements in the job description but absent from the master; never insert gaps into resume text.
- Preserve all bullet indices. Order arrays must be complete permutations of the supplied indices.
- Return only the requested schema-constrained JSON.

ACCEPTANCE GATE
- A plan that only reorders skills or bullets is invalid.
- Substantially rewrite CURRENT SUMMARY while preserving its facts and length limit.
- Surface at least {required_summary_skills} skill(s) from SUMMARY SKILLS TO SURFACE in the rewritten summary.
- Substantively rewrite at least one relevant experience or project bullet using only that bullet's facts.
- Keep every original number, tool and outcome attached to its original indexed bullet.
- Do not return the current summary verbatim, merely rearrange its words, or copy a bullet unchanged.
- Prefer concise ATS language from the job description only when exact master-resume evidence supports it.

JOB
Company: {context.company}
Title: {context.title}
Location: {context.location}
Job ID: {context.requisition_id}
URL: {context.url}
Description:
{description or "No detailed job description was available. Make only minimal, fully grounded changes."}

ALLOWED SKILLS
{json.dumps(template.skill_catalog, ensure_ascii=False)}

JD-SUPPORTED ALLOWED SKILLS
{json.dumps(supported_skills, ensure_ascii=False)}

SUMMARY SKILLS TO SURFACE
{json.dumps(summary_skills, ensure_ascii=False)}

CURRENT SUMMARY
{template.paragraphs[template.summary_index]}

EXPERIENCE BULLETS
{_indexed_bullets(template, template.experience_indices)}

PROJECT BULLETS
{_indexed_bullets(template, template.project_indices)}

MASTER RESUME FACTS
{template.full_text}
"""


def _response_json(response: requests.Response) -> dict[str, Any]:
    if not response.ok:
        message = ""
        try:
            error = response.json().get("error", {})
            if isinstance(error, dict):
                code = str(error.get("code", "") or "").strip()
                detail = str(error.get("message", "") or "").strip()
                message = ": ".join(value for value in (code, detail) if value)
        except (ValueError, AttributeError):
            pass
        raise ResumeTailoringError(
            f"Gemini API HTTP {response.status_code}"
            + (f" ({message[:500]})" if message else "")
        )
    payload = response.json()
    status = str(payload.get("status", "") or "").casefold()
    if status and status != "completed":
        raise ResumeTailoringError(
            f"Gemini stopped without a usable plan: {status}"
        )
    text_blocks: list[str] = []
    for step in payload.get("steps", []):
        if not isinstance(step, dict) or step.get("type") != "model_output":
            continue
        for block in step.get("content", []):
            if isinstance(block, dict) and block.get("type") == "text":
                text_blocks.append(str(block.get("text", "")))
    if not text_blocks and isinstance(payload.get("output_text"), str):
        text_blocks.append(payload["output_text"])
    if not text_blocks:
        raise ResumeTailoringError("Gemini returned no structured text")
    result = json.loads("".join(text_blocks))
    if not isinstance(result, dict):
        raise ResumeTailoringError("Gemini response was not a JSON object")
    return result


def _call_model(model: str, prompt: str, api_key: str) -> dict[str, Any]:
    response = requests.post(
        GEMINI_INTERACTIONS_URL,
        headers={
            "x-goog-api-key": api_key,
            "content-type": "application/json",
        },
        json={
            "model": model,
            "system_instruction": (
                "You are a truthful resume editor. Optimize relevance without "
                "adding any claim not explicitly supported by exact evidence from "
                "the supplied master resume. Return only schema-valid JSON."
            ),
            "input": prompt,
            "response_format": {
                "type": "text",
                "mime_type": "application/json",
                "schema": TAILOR_SCHEMA,
            },
        },
        timeout=(15, 120),
    )
    if response.status_code == 429:
        try:
            delay = min(15.0, max(0.0, float(response.headers.get("Retry-After", "0"))))
        except ValueError:
            delay = 0.0
        if delay:
            time.sleep(delay)
    return _response_json(response)


def request_tailoring_plan(
    template: TemplateSnapshot,
    job: Any,
    *,
    api_key: str | None = None,
) -> tuple[dict[str, Any], str, tuple[str, ...]]:
    key = (api_key if api_key is not None else os.getenv("GEMINI_API_KEY", "")).strip()
    if not key:
        raise ResumeTailoringError("GEMINI_API_KEY is missing")
    prompt = build_prompt(template, job)
    warnings: list[str] = []
    models = tuple(dict.fromkeys(model for model in (PRIMARY_MODEL, FALLBACK_MODEL) if model))
    for model in models:
        try:
            return _call_model(model, prompt, key), model, tuple(warnings)
        except Exception as exc:
            detail = str(exc).strip() or type(exc).__name__
            warnings.append(f"{model} failed: {detail[:600]}")
    raise ResumeTailoringError("; ".join(warnings) or "No Gemini model configured")


def _numbers(value: str) -> set[str]:
    return set(re.findall(r"(?<![A-Za-z])\d+(?:\.\d+)?%?\+?", value))


def _tokens(value: str) -> set[str]:
    return {
        token
        for token in re.findall(r"[a-z][a-z0-9+#./-]{2,}", value.casefold())
        if token not in STOPWORDS
    }


def _semantic_signature(value: str) -> tuple[tuple[str, int], ...]:
    """Ignore punctuation and word order when detecting substantive edits."""
    tokens = [
        token
        for token in re.findall(r"[a-z][a-z0-9+#./-]{2,}", value.casefold())
        if token not in STOPWORDS
    ]
    return tuple(sorted(Counter(tokens).items()))


def _evidence_is_valid(evidence: Any, master_text: str) -> bool:
    if not isinstance(evidence, list) or not evidence:
        return False
    normalized_master = _normalize(master_text)
    return all(
        isinstance(quote, str)
        and len(quote.strip()) >= 4
        and _normalize(quote) in normalized_master
        for quote in evidence
    )


def _claim_is_grounded(
    text: Any,
    evidence: Any,
    template: TemplateSnapshot,
    *,
    max_length: int,
    allowed_evidence_text: str | None = None,
) -> bool:
    if not isinstance(text, str) or not text.strip() or len(text) > max_length:
        return False
    lowered = text.casefold()
    if any(marker in lowered for marker in PLACEHOLDER_MARKERS):
        return False
    if any(marker in text for marker in ("\u200b", "\u200c", "\u200d")):
        return False
    evidence_source = allowed_evidence_text or template.full_text
    if not _evidence_is_valid(evidence, evidence_source):
        return False
    evidence_text = " ".join(str(value) for value in evidence)
    if not _numbers(text).issubset(_numbers(evidence_text)):
        return False
    for term in KNOWN_TECH_TERMS:
        if _contains_phrase(text, term) and not _contains_phrase(evidence_text, term):
            return False
    proposal_tokens = _tokens(text)
    if proposal_tokens:
        supported_tokens = _tokens(evidence_text)
        novel_ratio = len(proposal_tokens - supported_tokens) / len(proposal_tokens)
        if novel_ratio > 0.35:
            return False
    return True


def _valid_permutation(value: Any, size: int) -> list[int]:
    expected = list(range(size))
    if (
        isinstance(value, list)
        and all(isinstance(item, int) for item in value)
        and sorted(value) == expected
    ):
        return list(value)
    return expected


def _skill_lookup(template: TemplateSnapshot) -> dict[str, str]:
    return {_normalize(skill): skill for skill in template.skill_catalog}


def _skills_in_job(template: TemplateSnapshot, description: str) -> list[str]:
    matches: list[str] = []
    for skill in template.skill_catalog:
        normalized_skill = _normalize(skill)
        if normalized_skill and (
            _contains_phrase(description, skill)
            or any(
                len(token) >= 3 and _contains_phrase(description, token)
                for token in normalized_skill.split()
            )
        ):
            matches.append(skill)
    return matches


def _relevance_order(
    template: TemplateSnapshot, indices: tuple[int, ...], description: str
) -> list[int]:
    """Rank existing truthful bullets by JD overlap without rewriting claims."""
    description_tokens = _tokens(description)
    supported_skills = _skills_in_job(template, description)
    scored: list[tuple[int, int, int]] = []
    for local_index, paragraph_index in enumerate(indices):
        text = template.paragraphs[paragraph_index]
        skill_score = sum(
            1 for skill in supported_skills if _contains_phrase(text, skill)
        )
        token_score = len(_tokens(text) & description_tokens)
        scored.append((-skill_score, -token_score, local_index))
    return [local_index for _, _, local_index in sorted(scored)]


def safe_plan(template: TemplateSnapshot, job: Any) -> dict[str, Any]:
    context = _job_context(job)
    supported = _skills_in_job(template, context.description)
    gaps = [
        gap
        for gap in GAP_CANDIDATES
        if _contains_phrase(context.description, gap)
        and not _contains_phrase(template.full_text, gap)
    ][:5]
    return {
        "professional_summary": {
            "text": template.paragraphs[template.summary_index],
            "evidence": [template.paragraphs[template.summary_index]],
        },
        "skill_priorities": supported,
        "experience_bullets": [],
        "experience_order": _relevance_order(
            template, template.experience_indices, context.description
        ),
        "project_bullets": [],
        "project_order": _relevance_order(
            template, template.project_indices, context.description
        ),
        "supported_skills": supported[:8],
        "important_gaps": gaps,
    }


def validate_plan(
    raw_plan: dict[str, Any], template: TemplateSnapshot, job: Any
) -> tuple[dict[str, Any], tuple[str, ...]]:
    fallback = safe_plan(template, job)
    warnings: list[str] = []
    result = dict(fallback)
    summary = raw_plan.get("professional_summary")
    original_summary = template.paragraphs[template.summary_index]
    summary_text = (
        str(summary.get("text", "")).strip()
        if isinstance(summary, dict)
        else ""
    )
    if isinstance(summary, dict) and _claim_is_grounded(
        summary_text, summary.get("evidence"), template,
        max_length=len(original_summary),
    ) and _semantic_signature(summary_text) != _semantic_signature(original_summary):
        result["professional_summary"] = {
            "text": summary_text,
            "evidence": list(summary["evidence"]),
        }
    else:
        warnings.append("unsafe summary proposal rejected (including unchanged text)")

    lookup = _skill_lookup(template)
    priorities: list[str] = []
    for value in raw_plan.get("skill_priorities", []):
        if isinstance(value, str) and _normalize(value) in lookup:
            canonical = lookup[_normalize(value)]
            if canonical not in priorities:
                priorities.append(canonical)
    result["skill_priorities"] = priorities or fallback["skill_priorities"]

    for key, indices in (
        ("experience_bullets", template.experience_indices),
        ("project_bullets", template.project_indices),
    ):
        accepted: list[dict[str, Any]] = []
        seen_indices: set[int] = set()
        proposals = raw_plan.get(key, [])
        if not isinstance(proposals, list):
            proposals = []
        for proposal in proposals:
            if not isinstance(proposal, dict) or not isinstance(proposal.get("index"), int):
                warnings.append(f"malformed {key} proposal rejected")
                continue
            local_index = proposal["index"]
            if local_index in seen_indices or not 0 <= local_index < len(indices):
                warnings.append(f"invalid {key} index rejected")
                continue
            source = template.paragraphs[indices[local_index]]
            proposed_text = str(proposal.get("text", "")).strip()
            if not proposed_text.startswith("•"):
                proposed_text = "•  " + proposed_text.lstrip("-• ")
            if _claim_is_grounded(
                proposed_text, proposal.get("evidence"), template,
                max_length=len(source),
                allowed_evidence_text=source,
            ) and _semantic_signature(proposed_text) != _semantic_signature(source):
                accepted.append({
                    "index": local_index,
                    "text": proposed_text,
                    "evidence": list(proposal["evidence"]),
                })
                seen_indices.add(local_index)
            else:
                warnings.append(
                    f"unsafe or unchanged {key}[{local_index}] proposal rejected"
                )
        result[key] = accepted

    # Deterministic evidence-only ranking is safer and more consistent than
    # accepting an arbitrary model order. It also ensures the fallback can
    # materially tailor a resume without inventing or rewriting any claim.
    result["experience_order"] = fallback["experience_order"]
    result["project_order"] = fallback["project_order"]

    context = _job_context(job)
    supported: list[str] = []
    for value in raw_plan.get("supported_skills", []):
        if not isinstance(value, str) or _normalize(value) not in lookup:
            continue
        canonical = lookup[_normalize(value)]
        normalized_skill = _normalize(canonical)
        if _contains_phrase(context.description, canonical) or any(
            len(token) >= 3 and _contains_phrase(context.description, token)
            for token in normalized_skill.split()
        ):
            if canonical not in supported:
                supported.append(canonical)
    result["supported_skills"] = (supported or fallback["supported_skills"])[:8]

    gaps: list[str] = []
    for value in raw_plan.get("important_gaps", []):
        if not isinstance(value, str) or not 2 <= len(value.strip()) <= 80:
            continue
        normalized_gap = _normalize(value)
        if (
            normalized_gap
            and _contains_phrase(context.description, value)
            and not _contains_phrase(template.full_text, value)
            and value.strip() not in gaps
        ):
            gaps.append(value.strip())
    result["important_gaps"] = (gaps or fallback["important_gaps"])[:5]
    return result, tuple(warnings)


def _copy_run_format(target: Any, source: Any) -> None:
    source_properties = (
        deepcopy(source._r.rPr) if source._r.rPr is not None else None
    )
    target_properties = target._r.rPr
    if target_properties is not None:
        target._r.remove(target_properties)
    if source_properties is not None:
        target._r.insert(0, source_properties)


def _text_chunks(text: str, bold_phrases: list[str]) -> list[tuple[str, bool]]:
    phrases = sorted(
        {phrase.strip() for phrase in bold_phrases if phrase.strip()},
        key=len, reverse=True,
    )
    if not phrases:
        return [(text, False)]
    pattern = re.compile("(" + "|".join(re.escape(value) for value in phrases) + ")", re.I)
    chunks: list[tuple[str, bool]] = []
    for part in pattern.split(text):
        if not part:
            continue
        is_bold = any(part.casefold() == value.casefold() for value in phrases)
        if chunks and chunks[-1][1] == is_bold:
            chunks[-1] = (chunks[-1][0] + part, is_bold)
        else:
            chunks.append((part, is_bold))
    return chunks


def _rewrite_runs(paragraph: Paragraph, text: str) -> None:
    existing = list(paragraph.runs)
    if not existing:
        paragraph.add_run(text)
        return
    bold_phrases = [run.text for run in existing if run.bold and run.text.strip()]
    regular_template = next((run for run in existing if not run.bold), existing[0])
    bold_template = next((run for run in existing if run.bold), regular_template)
    chunks = _text_chunks(text, bold_phrases)
    while len(existing) < len(chunks):
        existing.append(paragraph.add_run())
    for index, run in enumerate(existing):
        if index < len(chunks):
            value, is_bold = chunks[index]
            _copy_run_format(run, bold_template if is_bold else regular_template)
            run.text = value
        else:
            run.text = ""


def _rewrite_skill_line(paragraph: Paragraph, priorities: list[str]) -> None:
    original = paragraph.text
    if ":" not in original:
        return
    label, values = original.split(":", 1)
    skills = [value.strip() for value in values.split(",") if value.strip()]
    rank = {_normalize(skill): index for index, skill in enumerate(priorities)}
    ordered = sorted(
        enumerate(skills),
        key=lambda item: (rank.get(_normalize(item[1]), len(rank)), item[0]),
    )
    content = ", ".join(skill for _, skill in ordered)
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(f"{label}:  {content}")
        return
    bold_template = next((run for run in runs if run.bold), runs[0])
    regular_template = next((run for run in runs if not run.bold), runs[-1])
    while len(runs) < 2:
        runs.append(paragraph.add_run())
    _copy_run_format(runs[0], bold_template)
    runs[0].text = f"{label}:  "
    _copy_run_format(runs[1], regular_template)
    runs[1].text = content
    for run in runs[2:]:
        run.text = ""


def _reorder_paragraphs(
    paragraphs: list[Paragraph], indices: tuple[int, ...], order: list[int]
) -> None:
    if not indices or order == list(range(len(indices))):
        return
    blocks = [paragraphs[index]._p for index in indices]
    anchor = paragraphs[max(indices) + 1]._p
    for local_index in order:
        anchor.addprevious(blocks[local_index])


def _sanitize_component(value: str, fallback: str, limit: int = 55) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_")
    return (cleaned or fallback)[:limit].rstrip("_")


def output_filename(job: Any) -> str:
    context = _job_context(job)
    job_id = context.requisition_id.strip()
    if not job_id:
        match = re.search(r"(?i)(?:jr|req|job)[-_ ]?\d{4,}|\d{6,}", context.url)
        job_id = match.group(0) if match else hashlib.sha256(
            context.url.encode("utf-8")
        ).hexdigest()[:10]
    return "_".join((
        _sanitize_component(context.company, "Company", 45),
        _sanitize_component(context.title, "Role", 55),
        _sanitize_component(job_id, "Job", 35),
        "Jagadev.docx",
    ))


def _restore_preserve_only_parts(master_path: Path, output_path: Path) -> None:
    """Keep every package part except the edited document XML byte-for-byte."""
    temporary_path = output_path.with_name(f".{output_path.name}.building")
    try:
        with (
            zipfile.ZipFile(master_path) as master_archive,
            zipfile.ZipFile(output_path) as generated_archive,
        ):
            generated_document_xml = generated_archive.read("word/document.xml")
            with zipfile.ZipFile(temporary_path, "w") as rebuilt:
                for item in master_archive.infolist():
                    content = (
                        generated_document_xml
                        if item.filename == "word/document.xml"
                        else master_archive.read(item.filename)
                    )
                    rebuilt.writestr(item, content)
        os.replace(temporary_path, output_path)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()


def render_tailored_resume(
    template: TemplateSnapshot,
    plan: dict[str, Any],
    job: Any,
    output_dir: Path | str,
) -> Path:
    destination_dir = Path(output_dir).resolve()
    destination_dir.mkdir(parents=True, exist_ok=True)
    output_path = destination_dir / output_filename(job)
    if output_path.resolve() == template.path.resolve():
        raise ResumeTailoringError("Refusing to overwrite the master resume")
    shutil.copy2(template.path, output_path)
    document = Document(output_path)
    paragraphs = list(document.paragraphs)
    _rewrite_runs(
        paragraphs[template.summary_index], plan["professional_summary"]["text"]
    )
    priorities = list(plan.get("skill_priorities", []))
    for index in template.skill_indices:
        _rewrite_skill_line(paragraphs[index], priorities)
    for key, indices in (
        ("experience_bullets", template.experience_indices),
        ("project_bullets", template.project_indices),
    ):
        for proposal in plan.get(key, []):
            _rewrite_runs(paragraphs[indices[proposal["index"]]], proposal["text"])
    _reorder_paragraphs(
        paragraphs, template.experience_indices, list(plan["experience_order"])
    )
    _reorder_paragraphs(
        paragraphs, template.project_indices, list(plan["project_order"])
    )
    # A shorter tailored section must not leave a heading alone at a page foot.
    for heading_index in template.headings.values():
        paragraphs[heading_index].paragraph_format.keep_with_next = True
    document.save(output_path)
    _restore_preserve_only_parts(template.path, output_path)
    return output_path


def validate_generated_resume(
    template: TemplateSnapshot, output_path: Path | str
) -> tuple[str, ...]:
    output = Path(output_path).resolve()
    if not output.is_file() or output.stat().st_size == 0:
        raise ResumeTailoringError("Generated resume is missing or empty")
    if _sha256(template.path) != template.sha256:
        raise ResumeTailoringError("Master resume changed during tailoring")
    if _package_entries(output) != template.package_entries:
        raise ResumeTailoringError("DOCX package structure changed unexpectedly")
    if _preserve_only_part_hashes(output) != _preserve_only_part_hashes(template.path):
        raise ResumeTailoringError("Generated resume changed preserve-only DOCX parts")
    if _hyperlink_targets(output) != template.hyperlink_targets:
        raise ResumeTailoringError("Generated resume hyperlinks changed")
    generated = Document(output)
    if len(generated.paragraphs) != len(template.paragraphs):
        raise ResumeTailoringError("Generated resume paragraph count changed")
    if _section_signature(generated) != template.section_signature:
        raise ResumeTailoringError("Generated resume page geometry changed")
    if tuple(
        "\n".join(paragraph.text for paragraph in section.header.paragraphs)
        for section in generated.sections
    ) != template.header_text:
        raise ResumeTailoringError("Generated resume header changed")
    if tuple(
        "\n".join(paragraph.text for paragraph in section.footer.paragraphs)
        for section in generated.sections
    ) != template.footer_text:
        raise ResumeTailoringError("Generated resume footer changed")
    text = "\n".join(paragraph.text for paragraph in generated.paragraphs)
    for required in (template.paragraphs[0], template.paragraphs[1], *REQUIRED_HEADINGS):
        if required not in text:
            raise ResumeTailoringError(f"Generated resume lost required content: {required}")
    lowered = text.casefold()
    if any(marker in lowered for marker in PLACEHOLDER_MARKERS):
        raise ResumeTailoringError("Generated resume contains a placeholder")
    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml")
    if b"w:vanish" in document_xml or b"w:delText" in document_xml:
        raise ResumeTailoringError("Generated resume contains hidden or deleted text")
    return ()


def _skill_hits(text: str, skills: list[str]) -> tuple[str, ...]:
    return tuple(skill for skill in skills if _contains_phrase(text, skill))


def _substantive_rewrite_count(
    original_values: tuple[str, ...], current_values: tuple[str, ...]
) -> int:
    """Count content edits while treating pure reordering as unchanged."""
    remaining = Counter(_semantic_signature(value) for value in original_values)
    rewritten = 0
    for value in current_values:
        signature = _semantic_signature(value)
        if remaining[signature]:
            remaining[signature] -= 1
        else:
            rewritten += 1
    return rewritten


def compare_ats_alignment(
    template: TemplateSnapshot, output_path: Path | str, job: Any
) -> ATSComparison:
    generated = Document(Path(output_path).resolve())
    current = tuple(paragraph.text for paragraph in generated.paragraphs)
    supported = _skills_in_job(template, _job_context(job).description)
    before_summary = template.paragraphs[template.summary_index]
    after_summary = current[template.summary_index]
    before_hits = _skill_hits(before_summary, supported)
    after_hits = _skill_hits(after_summary, supported)
    before_normalized = {_normalize(value) for value in before_hits}
    added = tuple(
        value for value in after_hits if _normalize(value) not in before_normalized
    )
    original_experience = tuple(
        template.paragraphs[index] for index in template.experience_indices
    )
    current_experience = tuple(current[index] for index in template.experience_indices)
    original_projects = tuple(
        template.paragraphs[index] for index in template.project_indices
    )
    current_projects = tuple(current[index] for index in template.project_indices)
    return ATSComparison(
        summary_keywords_before=before_hits,
        summary_keywords_after=after_hits,
        newly_surfaced_summary_keywords=added,
        experience_bullets_rewritten=_substantive_rewrite_count(
            original_experience, current_experience
        ),
        project_bullets_rewritten=_substantive_rewrite_count(
            original_projects, current_projects
        ),
    )


def validate_meaningful_ats_improvement(
    template: TemplateSnapshot,
    output_path: Path | str,
    job: Any,
    comparison: ATSComparison,
) -> None:
    generated = Document(Path(output_path).resolve())
    current_summary = generated.paragraphs[template.summary_index].text
    original_summary = template.paragraphs[template.summary_index]
    if _semantic_signature(current_summary) == _semantic_signature(original_summary):
        raise ResumeTailoringError(
            "Tailoring made no substantive professional-summary change"
        )
    supported = _skills_in_job(template, _job_context(job).description)
    before = {_normalize(value) for value in comparison.summary_keywords_before}
    available = [value for value in supported if _normalize(value) not in before]
    required_added = min(2, len(available))
    if len(comparison.newly_surfaced_summary_keywords) < required_added:
        raise ResumeTailoringError(
            "Tailoring did not surface enough JD-supported skills in the summary "
            f"({len(comparison.newly_surfaced_summary_keywords)}/{required_added})"
        )
    if comparison.rewritten_bullet_count < 1:
        raise ResumeTailoringError(
            "Tailoring only reordered existing content; at least one evidence-backed "
            "bullet must be substantively rewritten"
        )


def material_changes(
    template: TemplateSnapshot,
    output_path: Path | str,
    comparison: ATSComparison | None = None,
) -> tuple[str, ...]:
    generated = Document(Path(output_path).resolve())
    current = tuple(paragraph.text for paragraph in generated.paragraphs)
    original = template.paragraphs
    report = comparison or ATSComparison((), (), (), 0, 0)
    changes: list[str] = []
    if _semantic_signature(current[template.summary_index]) != _semantic_signature(
        original[template.summary_index]
    ):
        changes.append("professional summary rewritten")
    if any(current[index] != original[index] for index in template.skill_indices):
        changes.append("skills prioritized")
    if report.experience_bullets_rewritten:
        changes.append(
            f"experience bullets rewritten ({report.experience_bullets_rewritten})"
        )
    elif tuple(current[index] for index in template.experience_indices) != tuple(
        original[index] for index in template.experience_indices
    ):
        changes.append("experience bullets reordered")
    if report.project_bullets_rewritten:
        changes.append(
            f"project bullets rewritten ({report.project_bullets_rewritten})"
        )
    elif tuple(current[index] for index in template.project_indices) != tuple(
        original[index] for index in template.project_indices
    ):
        changes.append("project bullets reordered")
    return tuple(changes)


def generate_tailored_resume(
    job: Any,
    *,
    master_path: Path | str | None = None,
    output_dir: Path | str | None = None,
    require_ai: bool = False,
    api_key: str | None = None,
) -> TailoredResume:
    master = Path(
        master_path or os.getenv("MASTER_RESUME_PATH", "") or DEFAULT_MASTER_RESUME
    )
    destination = Path(
        output_dir or os.getenv("TAILORED_RESUME_DIR", "") or DEFAULT_OUTPUT_DIR
    )
    template = inspect_template(master)
    warnings: list[str] = []
    model = "safe-template"
    try:
        raw_plan, model, request_warnings = request_tailoring_plan(
            template, job, api_key=api_key
        )
        warnings.extend(request_warnings)
        plan, validation_warnings = validate_plan(raw_plan, template, job)
        warnings.extend(validation_warnings)
    except Exception as exc:
        if require_ai:
            raise ResumeTailoringError(str(exc)) from exc
        warnings.append(f"AI tailoring unavailable: {type(exc).__name__}")
        plan = safe_plan(template, job)
    try:
        output_path = render_tailored_resume(template, plan, job, destination)
        validate_generated_resume(template, output_path)
        comparison = compare_ats_alignment(template, output_path, job)
        validate_meaningful_ats_improvement(
            template, output_path, job, comparison
        )
        changed_sections = material_changes(template, output_path, comparison)
        if not changed_sections:
            raise ResumeTailoringError(
                "Tailoring produced no material change; refusing unchanged attachment"
            )
    except Exception as exc:
        raise ResumeTailoringError(str(exc)) from exc
    return TailoredResume(
        path=output_path,
        supported_skills=tuple(plan.get("supported_skills", ()))[:8],
        important_gaps=tuple(plan.get("important_gaps", ()))[:5],
        model=model,
        warnings=tuple(warnings),
        changed_sections=changed_sections,
        comparison=comparison,
        source_path=template.path,
    )
